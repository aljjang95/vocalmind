"""
강한소상공인 성장지원 사업 신청서 — 보컬마인드 .docx 생성
양식: 붙임 2 강한소상공인 성장지원 사업 신청서식(사업계획서 양식)
출력: C:/Users/Administrator/Desktop/보컬마인드_강한소상공인_사업계획서.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(r"C:/Users/Administrator/Desktop/보컬마인드_강한소상공인_사업계획서.docx")


def set_cell_bg(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "555555")
        borders.append(b)
    tc_pr.append(borders)


def write_run(p, text: str, bold: bool = False, size: int = 10, color: str | None = None) -> None:
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 16, 2: 13, 3: 11}
    write_run(p, text, bold=True, size=sizes.get(level, 11), color="1F3A5F")


def add_body(doc: Document, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    write_run(p, text, bold=bold, size=size)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    write_run(p, "• " + text, size=10)


def add_table_kv(doc: Document, rows: list[tuple[str, str]], key_w: float = 4.5) -> None:
    """양식 표준: 좌측 항목명(짙은 회색) + 우측 내용."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(key_w)
    tbl.columns[1].width = Cm(13.5)
    for i, (k, v) in enumerate(rows):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        c0.width = Cm(key_w)
        c1.width = Cm(13.5)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_bg(c0, "E8EEF5")
        set_cell_borders(c0)
        set_cell_borders(c1)
        c0.text = ""
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        write_run(p0, k, bold=True, size=10, color="1F3A5F")
        c1.text = ""
        for j, line in enumerate(v.split("\n")):
            p1 = c1.paragraphs[0] if j == 0 else c1.add_paragraph()
            p1.paragraph_format.space_after = Pt(2)
            write_run(p1, line, size=10)


def add_grid_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    for i, w in enumerate(widths_cm):
        tbl.columns[i].width = Cm(w)
    # header
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        c.width = Cm(widths_cm[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(c, "1F3A5F")
        set_cell_borders(c)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        write_run(p, h, bold=True, size=10, color="FFFFFF")
    # body
    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.cell(r, i)
            c.width = Cm(widths_cm[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(c)
            c.text = ""
            for j, line in enumerate(val.split("\n")):
                p = c.paragraphs[0] if j == 0 else c.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                write_run(p, line, size=9)


def add_section_break(doc: Document) -> None:
    p = doc.add_paragraph()
    write_run(p, "─" * 60, size=8, color="888888")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ───────────── 표지 ─────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    write_run(p, "2026년 강한소상공인 성장지원 사업", bold=True, size=20, color="1F3A5F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    write_run(p, "사업계획서", bold=True, size=24, color="1F3A5F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    write_run(p, "AI가 발성을 진단하고 본인 MR로 뮤직비디오까지 만들어주는", size=12, color="333333")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(p, "1인 보컬 스튜디오 「보컬마인드」", bold=True, size=14, color="1F3A5F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    write_run(p, "신 청 자 :  비너스(1인 대표)", size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(p, "사 업 명 :  HLB 보컬스튜디오 (보컬마인드)", size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    write_run(p, "제 출 일 :  2026년 4월", size=12)

    doc.add_page_break()

    # ───────────── 1. 한 줄 소개 ─────────────
    add_heading(doc, "1. 제품(서비스) 한 줄 소개", 1)
    add_body(doc,
             "AI가 발성을 진단하고 본인 MR로 뮤직비디오까지 만들어주는 1인 보컬 스튜디오 「보컬마인드」",
             bold=True, size=12)
    add_body(doc, "(보조안 — 분량 조정 시 택일)", size=9)
    add_bullet(doc, "AI 발성 코치 + 본인 음색 커버 MV 생성기 「HLB 보컬마인드」")
    add_bullet(doc, "노래방 무대의 주인공이 되는 AI 보컬 트레이닝 「보컬마인드」")
    add_section_break(doc)

    # ───────────── 2. 기업 소개 ─────────────
    add_heading(doc, "2. 기업 소개", 1)

    add_heading(doc, "지원동기 및 본 사업 참가 취지", 2)
    add_body(doc,
             "보컬 학원은 회당 5만~15만 원, 한 달이면 30만 원 이상이 든다. 비싸게 다녀도 강사가 "
             "「혀에 힘 빼세요」, 「후두 내리세요」라고 말해줄 뿐, 본인이 그게 정말 됐는지 체감으로 "
             "확인할 방법이 없었다. 이 문제를 AI로 풀고 싶었다.")
    add_body(doc,
             "발성의 4축(후두·혀뿌리·턱·성구전환) 긴장을 음성 신호에서 직접 측정해 숫자와 시각화로 "
             "보여주고, 누구든 혼자서 발성을 교정할 수 있도록 만들었다.")
    add_body(doc,
             "이미 1인 풀스택으로 백엔드(379개 자동화 테스트), 프론트엔드, GPU 워커 3종, 결제, "
             "AI 스튜디오까지 전부 구축 완료했고 베타 진입 직전이다. 본 사업의 성장지원자금은 "
             "(1) Modal GPU·Runware API 비용 풀, (2) 1인 보컬강사 파트너 매칭 인센티브, "
             "(3) 글로벌 진출 다국어 작업, 세 영역에 집중 투입한다.")

    add_heading(doc, "아이템 및 수행 역량", 2)
    add_grid_table(doc,
                   headers=["영역", "구축 현황"],
                   rows=[
                       ["AI 발성 진단 엔진",
                        "parselmouth(Praat 기반 음성 분석 라이브러리) 4축 긴장 측정\n"
                        "Jitter / Shimmer / HNR / 포먼트(F1·F2) / VSA / 성구 smoothness"],
                       ["AI 코칭",
                        "Anthropic Claude Haiku + ChromaDB 벡터 DB\n"
                        "코칭 데이터 1,850건 + 커리큘럼 211건 임베딩 완료"],
                       ["AI 스튜디오",
                        "본인 MR → Demucs 분리 → RVC(음색 클로닝) → FLUX(이미지) → "
                        "HunyuanVideo(텐센트 영상 생성) → LatentSync(립싱크) → FFmpeg 합성 → "
                        "C2PA(AI 콘텐츠 표준 워터마크)"],
                       ["HLB 28단계 커리큘럼",
                        "발성 기초~믹스보이스까지 5-Phase 레슨 구조\n"
                        "(왜→시범→실습→평가→요약), 시범 영상 23단계 매핑 완료"],
                       ["결제 인프라",
                        "토스페이먼츠 선불 크레딧 (1편 5,000원 / 마진 60%)\n"
                        "멱등성 보장 + 자동 환불 RPC"],
                       ["배포 인프라",
                        "Vercel(프론트) + Fly.io(백엔드) + Modal Labs T4 GPU 3개 앱 + "
                        "Supabase Auth/DB/Storage"],
                       ["보안·신뢰",
                        "server-only API / RLS 7개 정책 / 3단계 모더레이션 / "
                        "C2PA 워터마크 (AI 영상 윤리 선제 대응)"],
                   ],
                   widths_cm=[4.0, 14.0])

    add_heading(doc, "차별성·혁신성", 2)
    add_bullet(doc, "국내 유일 통합 — 발성 4축 정량 진단 + AI 감각 코칭 + 본인 음색 커버 MV를 한 자리에서 제공")
    add_bullet(doc, "저작권 안전 — 본인 MR만 받기 때문에 가요 라이선스 리스크 0")
    add_bullet(doc, "종량제 결제 — 월정액 부담 없이 1편 5,000원, 필요할 때만")
    add_bullet(doc, "데이터 자산화 — 사용자가 늘수록 RAG 코칭 정확도가 올라가는 자기강화 루프")
    add_bullet(doc, "AI 영상 윤리 선제 대응 — C2PA 표준 워터마크로 딥페이크 우려 차단")

    add_heading(doc, "소비 트렌드 적합성", 2)
    add_bullet(doc, "MZ 자존감 콘텐츠 시대 — 「내 노래를 내 무대로」 수요 폭증")
    add_bullet(doc, "글로벌 K-POP 팬덤 — 본인 커버 MV 글로벌 공유 욕구")
    add_bullet(doc, "AI 영상 생성(Runway·Sora·Veo) 보편화 → 일반인도 콘텐츠 제작자")

    add_heading(doc, "대표자 및 인원 현황", 2)
    add_table_kv(doc, [
        ("대표(1인)",
         "풀스택 AI 도구 활용으로 통상 5인 팀 분량을 1인이 운영"),
        ("운영 도구",
         "Claude Code(Anthropic) / Cursor / Modal / Vercel / Supabase / 토스페이먼츠"),
        ("선행 사업 운영",
         "clauseai.kr (법률·부동산 자동 콘텐츠 플랫폼) — 1인 풀스택 운영 검증 완료"),
        ("추진 중 협업",
         "1인 보컬강사 / 실용음악 학원 1곳 사전 협의 진행 중"),
    ])
    add_section_break(doc)

    # ───────────── 3. 사업내용 ─────────────
    add_heading(doc, "3. 사업내용 (제품/서비스) — 시장성·경쟁력", 1)

    add_heading(doc, "시장성 — 글로벌 시장 규모 (객관 데이터)", 2)
    add_grid_table(doc,
                   headers=["구분", "규모", "출처"],
                   rows=[
                       ["음악 AI 시장 (2030)", "약 27억 달러 (≈ 4조 원)", "Forbes Korea"],
                       ["음성·음성인식 시장 (2026)", "237억 달러", "Fortune Business Insights"],
                       ["음성·음성인식 시장 (2034)", "1,040억 달러", "Fortune Business Insights"],
                       ["아시아·태평양 음성 시장 (2026)", "53.7억 달러 (최고 성장률 권역)", "Fortune Business Insights"],
                   ],
                   widths_cm=[5.5, 6.5, 6.0])

    add_heading(doc, "국내 보컬 시장", 2)
    add_bullet(doc, "실용음악·보컬 학원 약 2,500곳 (월 수강료 평균 20~30만 원)")
    add_bullet(doc, "노래방 시장 연 1조 원대 (1인 코인노래방 회복기)")
    add_bullet(doc, "잠재 수요 — 학원 다닐 시간·돈 부담스러운 20~40대 취미 보컬 + MZ 콘텐츠 크리에이터")

    add_heading(doc, "타겟 고객 (3계층)", 2)
    add_grid_table(doc,
                   headers=["계층", "타겟", "결제 시나리오"],
                   rows=[
                       ["1차 (B2C)", "노래 잘하고 싶은데 학원 부담스러운 20~40대",
                        "월 1~3편 × 5,000원"],
                       ["2차 (B2C)", "본인 커버 영상 SNS·유튜브 업로드 MZ 크리에이터",
                        "월 5~10편 + Voice Identity 등록"],
                       ["3차 (B2B)", "1인 보컬강사 + 실용음악 학원",
                        "AI 진단 매칭 수수료 + B2B 라이선스"],
                   ],
                   widths_cm=[3.0, 9.0, 6.0])

    add_heading(doc, "경쟁력 — 경쟁사 대비 우위", 2)
    add_grid_table(doc,
                   headers=["구분", "Smule", "오토튠 등 보정 SW", "일반 보컬 학원", "보컬마인드"],
                   rows=[
                       ["발성 진단", "음정만", "진단 X", "강사 주관", "4축 긴장 정량 측정"],
                       ["코칭", "점수 표시", "X", "강사 1:1", "AI 감각 코칭 + 28단계 커리큘럼"],
                       ["결과물", "녹음 파일", "보정된 음원", "수업 후기", "본인 음색 + MV 영상"],
                       ["저작권", "라이선스 의존", "위험", "해당 없음", "본인 MR만 (안전)"],
                       ["가격", "월 정액", "단발", "월 30만원+", "5,000원/편"],
                       ["글로벌 가능성", "있음", "있음", "없음", "있음"],
                   ],
                   widths_cm=[3.0, 3.0, 3.5, 3.0, 5.5])

    add_heading(doc, "사업 흐름 (Value Chain)", 2)
    flow = (
        "[사용자] 회원가입 + Voice Identity 10문장 등록 (한 번만)\n"
        "    ↓\n"
        "[사용자] 본인 MR 업로드 + 가이드 보컬 녹음\n"
        "    ↓\n"
        "[AI 진단] parselmouth 4축 긴장 측정 + Claude RAG 코칭\n"
        "    ↓\n"
        "[AI 스튜디오] Modal GPU(Demucs/RVC) → Runware(FLUX/HunyuanVideo/LatentSync) → "
        "Modal GPU(FFmpeg 합성)\n"
        "    ↓\n"
        "[모더레이션 + C2PA 워터마크]\n"
        "    ↓\n"
        "[사용자] 16:9 / 9:16 MP4 다운로드\n"
        "    ↓\n"
        "[선택] 1인 보컬강사 매칭 → 진단 데이터 공유 → 본 수업 (수수료 20%)"
    )
    add_body(doc, flow, size=10)

    add_heading(doc, "이해관계자", 2)
    add_grid_table(doc,
                   headers=["주체", "역할", "가치 교환"],
                   rows=[
                       ["사용자(B2C)", "발성 진단·연습·MV 제작", "5,000원/편 결제"],
                       ["1인 보컬강사", "AI 진단 후 본 수업", "매칭 학생 + 수수료 분배"],
                       ["실용음악 학원", "B2B 학생 진단 도구", "월 30만 원 라이선스"],
                       ["결제 인프라", "토스페이먼츠", "결제 수수료"],
                       ["AI 인프라", "Anthropic·Runware·Modal·Supabase", "API 비용 (원가)"],
                       ["콘텐츠 크리에이터", "본인 커버 MV SNS 업로드", "자연 유입 루프"],
                   ],
                   widths_cm=[4.0, 7.0, 7.0])
    add_section_break(doc)

    # ───────────── 4. 사업성 (혁신) ─────────────
    add_heading(doc, "4. 사업성 (혁신·성공전략)", 1)

    add_heading(doc, "우리만의 사업성", 2)
    add_table_kv(doc, [
        ("데이터 해자(moat)",
         "사용자 발성 진단 데이터가 누적될수록 RAG 코칭 정확도가 올라가는 자기강화 구조.\n"
         "현재 ChromaDB 1,850건 임베딩. 후발 진입자는 같은 양의 한국어 보컬 데이터를 모아야 함."),
        ("종량제 마진 구조",
         "1편당 원가 약 2,000원 / 판매가 5,000원 / 마진 60%.\n"
         "사용자 0이어도 손실 0(서버리스 + auto_stop), 사용자 폭증해도 인프라 자동 스케일."),
        ("Voice Identity 영구 자산",
         "사용자가 한 번 자기 음색을 등록하면 영구 사용 → 재구매 동기 강함, LTV 상승."),
        ("AI 영상 윤리 선제 대응",
         "정부의 AI 콘텐츠 표시 의무화 추세에 미리 대응 (C2PA 표준 워터마크).\n"
         "공공기관·미디어 협업 가능성 확보."),
    ], key_w=4.0)

    add_heading(doc, "혁신 아이디어 — 성공전략", 2)
    add_table_kv(doc, [
        ("전략 A. AI 진단 → 강사 추천 역방향 매칭",
         "일반 학원은 학생을 모집하고 진단하지만, 우리는 AI가 먼저 진단하고 필요한 학생만 적합한 "
         "강사에게 보낸다. 강사는 영업 비용 0, 학생은 본인 약점 알고 시작, 보컬마인드는 매칭 "
         "수수료 + 양 측 락인."),
        ("전략 B. 자존감 콘텐츠 자연 유입 루프",
         "사용자가 본인 커버 MV를 SNS에 공유 → 신규 유저 유입 → CAC 0원.\n"
         "이미 보유한 글로벌 SNS 자동 업로드 인프라를 Phase 1에 즉시 연결 가능."),
        ("전략 C. 글로벌 K-POP 보컬 흐름 흡수",
         "parselmouth는 음성 신호 분석이라 언어 무관. UI 다국어화만 하면 K-POP 글로벌 팬이 "
         "한국어 곡을 부르며 발성을 교정받는 동선 확보."),
        ("전략 D. 지역자원 기반 활성화",
         "지역 1인 보컬강사·실용음악 학원과 협업 → 지방 학생도 수도권급 코칭 접근.\n"
         "지역 노래방 협업 가능 (코인노래방 + 보컬마인드 진단 부스 융합)."),
    ], key_w=5.5)
    add_section_break(doc)

    # ───────────── 5. 주요 성과 ─────────────
    add_heading(doc, "5. 주요 성과", 1)

    add_heading(doc, "기술 성과 (개발 완료된 자산)", 2)
    add_bullet(doc, "백엔드 자동화 테스트 379개 PASS (parselmouth + Claude + Runware + 결제 + 모더레이션)")
    add_bullet(doc, "Phase 0 Bootstrap 완료 — Demucs/RVC/Compose 3개 Modal GPU 워커 + 13단계 상태머신 + 자동 환불 RPC")
    add_bullet(doc, "HLB 28단계 커리큘럼 데이터화 완료 (시범 오디오 + 23단계 YouTube 영상 구간 매핑)")
    add_bullet(doc, "AI 코칭 RAG — vocal_curriculum 211건 + vocal_feedback 1,850건 ChromaDB 임베딩")
    add_bullet(doc, "결제 시스템 — 토스페이먼츠 위젯 + 멱등성 + 자동 환불 + 선불 크레딧 잔액 view")

    add_heading(doc, "인프라 성과", 2)
    add_bullet(doc, "배포 완료 — Vercel(vocalmind-lemon.vercel.app) + Fly.io(nrt) + Modal 3개 앱")
    add_bullet(doc, "Supabase — Auth + 5개 테이블 + 4개 Storage 버킷 + RLS 7개 정책 + Realtime publication")
    add_bullet(doc, "Runware API 라이브 검증 완료 — FLUX Schnell 512×512 정상 호출")
    add_bullet(doc, "C2PA 워터마크 — 서명 키/인증서 박제 완료, AI 영상 표준 식별")

    add_heading(doc, "보안 성과", 2)
    add_bullet(doc, "server-only Anthropic 강제 (API 키 노출 차단)")
    add_bullet(doc, "3단계 모더레이션 (업로드 / Voice Identity / MV 출력) + moderation_events 감사 로그")
    add_bullet(doc, "X-Orchestrator-Secret 64hex 헤더 인증 (Modal ↔ FastAPI)")

    add_heading(doc, "지재권 / 입상 (추진 중)", 2)
    add_bullet(doc, "특허 출원 검토 중 — 한국어 발성 4축 긴장 측정 알고리즘 (parselmouth 응용)")
    add_bullet(doc, "상표 출원 예정 — 「보컬마인드」, 「HLB 보컬스튜디오」")
    add_bullet(doc, "선행 사업 운영 경험 — clauseai.kr (법률·부동산 자동 콘텐츠 플랫폼) 1인 풀스택 운영")

    add_heading(doc, "운영 KPI (베타 직전 목표)", 2)
    add_bullet(doc, "베타 1개월 — 가입 100명 / 결제 전환 20명 / 첫 매출 100만 원")
    add_bullet(doc, "보수적 추정 — 단가 5,000원 × 200건 = 100만 원")
    add_section_break(doc)

    # ───────────── 6. 향후 사업계획 ─────────────
    add_heading(doc, "6. 향후 사업계획 (성장성)", 1)

    add_heading(doc, "단기 (3개월) — Phase 0 베타 검증", 2)
    add_bullet(doc, "베타 오픈 → 가입 100명 / 결제 전환 20명 / 매출 100만 원")
    add_bullet(doc, "검증 KPI — D7 재방문율 ≥ 30% / 1편당 GPU 원가 ≤ 2,000원 / NPS ≥ 40")
    add_bullet(doc, "사용자 인터뷰 20건 → Phase 1 우선순위 확정")

    add_heading(doc, "중기 (4~9개월) — Phase 1 확장", 2)
    add_grid_table(doc,
                   headers=["항목", "내용"],
                   rows=[
                       ["정기구독 도입", "취미반 100,000원 / 발성전문반 150,000원 (월 정액)"],
                       ["1인 보컬강사 파트너 정식 오픈", "강사 50명 모집 / 매칭 수수료 20%"],
                       ["SNS 자동 업로드", "TikTok / Instagram / YouTube Shorts (글로벌 쿠키 매니저 활용)"],
                       ["목표 매출", "월 1,000만 원"],
                       ["누적 사용자", "5,000명"],
                   ],
                   widths_cm=[5.0, 13.0])

    add_heading(doc, "장기 (10~24개월) — Phase 2 글로벌 + B2B", 2)
    add_grid_table(doc,
                   headers=["항목", "내용"],
                   rows=[
                       ["B2B 학원 플랜", "실용음악 학원 100곳 도입 (월 30만 원)"],
                       ["글로벌 진출", "영어/일본어 UI → K-POP 보컬 글로벌 수요 흡수"],
                       ["AI 음성 데이터 라이선싱", "의료(음성장애)·교육(언어 발음) 분야 B2B"],
                       ["목표 매출", "월 5,000만 원"],
                       ["누적 사용자", "20,000명"],
                   ],
                   widths_cm=[5.0, 13.0])

    add_heading(doc, "홍보·마케팅 계획", 2)
    add_bullet(doc, "자연 유입 루프 (CAC 0원) — 본인 커버 MV → SNS 공유 → 신규 가입")
    add_bullet(doc, "유튜브 보컬 인플루언서 협찬 5명 × 3개월 (월 200만 원)")
    add_bullet(doc, "보컬강사 인플루언서 파트너십 — 강사 1인당 학생 매칭 수수료 분배")
    add_bullet(doc, "자체 SNS 자동화 인프라 활용 — 일 N건 숏츠 자동 생성·업로드")

    add_heading(doc, "자본 조달 계획", 2)
    add_bullet(doc, "현재 — 자기자본 운영 (월 50만 원 미만 / Vercel·Fly.io·Modal·Anthropic 비용)")
    add_bullet(doc, "단기 — 본 강한소상공인 사업 성장지원자금 + 사업화 자금 (1~2단계)")
    add_bullet(doc, "중기 — Pre-A 시드 라운드 (베타 결과 검증 후)")

    add_heading(doc, "1억 원 자금 사용 계획 (3단계 통과 가정)", 2)
    add_grid_table(doc,
                   headers=["항목", "비중", "금액", "사유"],
                   rows=[
                       ["Modal GPU + Runware API 비용 풀", "30%", "3,000만 원",
                        "사용자 폭증 대비 종량제 원가 선매입"],
                       ["1인 보컬강사 파트너 초기 인센티브", "20%", "2,000만 원",
                        "강사 50명 모집·매칭 수수료 면제 풀"],
                       ["인플루언서 마케팅", "20%", "2,000만 원",
                        "유튜브 보컬 인플루언서 5명 × 3개월"],
                       ["다국어화 (영어·일본어) + 글로벌 결제", "15%", "1,500만 원",
                        "K-POP 글로벌 진출 인프라"],
                       ["특허·상표·법무", "10%", "1,000만 원",
                        "발성 진단 알고리즘 특허 + 상표"],
                       ["운영 안전망 (인건비 1인 추가)", "5%", "500만 원",
                        "보컬 도메인 전문가 채용 검토"],
                   ],
                   widths_cm=[6.0, 1.8, 3.0, 7.2])

    add_heading(doc, "투자유치 계획", 2)
    add_bullet(doc, "베타 결과 검증 (가입 5,000명 / 월 매출 1,000만 원) 확보 시 Pre-A 라운드 검토")
    add_bullet(doc, "투자유치 성공 시 — 강사 매칭 알고리즘 고도화 + 글로벌 진출 가속 + 채용 확대")
    add_section_break(doc)

    # ───────────── 7. 기대효과 ─────────────
    add_heading(doc, "7. 기대효과 — 성과 창출 가능성", 1)

    add_heading(doc, "정량 효과 (12개월 기준)", 2)
    add_bullet(doc, "누적 사용자 20,000명")
    add_bullet(doc, "월 매출 3,000만 원 (단건 결제 + 정기구독 + B2B 학원)")
    add_bullet(doc, "1인 보컬강사 파트너 50명 / B2B 학원 10곳")

    add_heading(doc, "정성 효과", 2)
    add_table_kv(doc, [
        ("소상공인 보컬강사 매출 증대",
         "1인 강사가 학생 모집·진단에 쓰던 시간 → 수업 자체에 투입.\n"
         "AI가 1차 진단 대행 → 강사는 본인 강점(피드백·교정)에 집중.\n"
         "강사 1인당 월 매출 30~50% 증가 추정."),
        ("K-콘텐츠 산업 저변 확대",
         "일반인 커버 MV 진입장벽 해소 → K-POP 팬덤·창작자 확대 → 한국 콘텐츠 산업 동반 성장."),
        ("보컬 사교육 비용 양극화 해소",
         "월 30만 원 학원이 부담스러운 학습자에게 1편 5,000원의 대안 제공.\n"
         "보컬 학습 접근성 평등화."),
        ("AI 발성 진단 데이터 자산화",
         "한국어 보컬 빅데이터 축적 → 향후 의료(음성장애 진단)·교육(언어 발음 교정) B2B 확장.\n"
         "단순 앱 → 데이터 인프라 기업 전환."),
        ("지역경제 활성화",
         "지방 1인 보컬강사·실용음악 학원도 보컬마인드 매칭으로 수도권급 학생 유입.\n"
         "지역 노래방 협업 → 코인노래방 + 보컬 진단 부스 융합 모델."),
    ], key_w=5.5)
    add_section_break(doc)

    # ───────────── 8. 파트너 협업 ─────────────
    add_heading(doc, "8. 파트너사와 협업 계획", 1)

    add_body(doc,
             "※ 강한소상공인 평가에서 「파트너사 협업을 통한 고도화」가 별도 평가 항목으로 명시됨. "
             "단독 신청 시에도 추후 매칭 후 협업 계획을 반드시 작성.",
             size=9)

    add_heading(doc, "1순위 파트너 후보", 2)
    add_body(doc, "1인 운영 실용음악·보컬 학원 / 프리랜서 보컬강사 (지역 우선)", bold=True)

    add_heading(doc, "협업 모델 (사업기간 내 실행)", 2)
    flow2 = (
        "[사용자] 보컬마인드 가입 → AI 1차 진단 → 발성 약점 데이터 생성\n"
        "    ↓\n"
        "[보컬마인드 매칭 엔진] 지역·전문분야·요금제로 강사 자동 매칭\n"
        "    ↓\n"
        "[강사] AI 진단 데이터 + 음성 클립 검토 → 사용자에게 본 수업 제안\n"
        "    ↓\n"
        "[사용자] 강사 수업 결제 → 보컬마인드 매칭 수수료 20% 수취\n"
        "    ↓\n"
        "[강사 대시보드] 학생 진도·진단 추이 모니터링 → 수업 효율 극대화"
    )
    add_body(doc, flow2, size=10)

    add_heading(doc, "사업기간 내 협업 추진 일정", 2)
    add_grid_table(doc,
                   headers=["시기", "내용"],
                   rows=[
                       ["M1", "강사 5명 사전 인터뷰 → 워크플로우 검증"],
                       ["M2", "강사 대시보드 MVP 개발 (Next.js + Supabase)"],
                       ["M3", "강사 50명 모집 캠페인 (인플루언서 협업)"],
                       ["M4", "매칭 베타 오픈 → 첫 매칭 100건"],
                       ["M5~6", "매칭 데이터 분석 → 알고리즘 고도화"],
                   ],
                   widths_cm=[2.0, 16.0])

    add_heading(doc, "성장지원자금 사용 계획 (파트너 협업 영역, 사업화자금과 별도)", 2)
    add_grid_table(doc,
                   headers=["항목", "금액", "사유"],
                   rows=[
                       ["강사 모집 마케팅", "800만 원", "보컬강사 인플루언서 협찬 (강사 50명 목표)"],
                       ["강사 인센티브 풀 (첫 5건 수수료 면제)", "700만 원", "강사 진입장벽 0 → 매칭 활성화"],
                       ["강사 대시보드 개발", "400만 원", "UI/UX + 진단 데이터 시각화"],
                       ["매칭 알고리즘 고도화", "300만 원", "지역·전문분야·요금제 자동 매칭"],
                       ["파트너 워크숍·컨퍼런스", "200만 원", "강사 50명 대상 보컬마인드 사용법 교육"],
                       ["합계", "2,400만 원", ""],
                   ],
                   widths_cm=[7.0, 3.0, 8.0])

    add_heading(doc, "협업에 따른 기대효과", 2)
    add_table_kv(doc, [
        ("강사 입장",
         "학생 모집 비용 0 (보컬마인드가 매칭) / 1차 진단 자동화 → 본 수업 효율 200%↑ / "
         "매월 안정적 매칭 학생 유입 → 매출 30~50%↑"),
        ("보컬마인드 입장",
         "사용자가 강사 추천을 통해 보컬마인드를 신뢰 → 리텐션·LTV 상승 / 매칭 수수료 신규 매출 라인 / "
         "강사 입소문 자연 유입 → CAC 하락"),
        ("사용자 입장",
         "본인 약점 알고 강사 만남 → 비용·시간 절감 / 학습 효율 상승 (AI 진단 + 강사 1:1 시너지) / "
         "지방 거주자도 수도권급 강사 매칭 가능"),
        ("삼방향 윈윈",
         "AI(진단) + 강사(전문) + 사용자(학습) 분업 구조로 모두 강점만 살림."),
    ], key_w=4.5)

    add_heading(doc, "지역경제 상생 전략", 2)
    add_bullet(doc, "지역 1인 강사 우선 매칭 — 사용자 거주지 기반 매칭 가중치")
    add_bullet(doc, "지역 학원 B2B 도입 — 학원 1차 진단 도구로 활용 → 학원 매출 ↑")
    add_bullet(doc, "지역 노래방 협업 — 코인노래방 + 보컬마인드 진단 부스 결합 모델 (Phase 2)")
    add_bullet(doc, "지역 인플루언서 강사 발굴 — 지방 도시별 1인 보컬강사 인플루언서 1명 + 학생 매칭")

    # ───────────── 마무리 ─────────────
    doc.add_page_break()
    add_heading(doc, "[ 끝 ]", 1)
    add_body(doc,
             "본 사업계획서는 보컬마인드(1인 대표 비너스)의 2026 강한소상공인 성장지원 사업 신청용입니다. "
             "신청 마감일·시각자료 첨부·특허 출원 접수번호 확보를 사전 확인 후 제출 부탁드립니다.",
             size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
